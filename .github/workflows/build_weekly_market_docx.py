from pathlib import Path
import json
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

PACKET_JSON_PATH = OUTPUT_DIR / "weekly_market_packet.json"
DOCX_OUTPUT_PATH = OUTPUT_DIR / "weekly_market_packet.docx"


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(14)

    if "Title" in styles:
        styles["Title"].font.name = "Calibri"
        styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        styles["Title"].font.size = Pt(22)
        styles["Title"].font.bold = True

    if "Heading 1" in styles:
        styles["Heading 1"].font.name = "Calibri"
        styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        styles["Heading 1"].font.size = Pt(16)
        styles["Heading 1"].font.bold = True

    if "Heading 2" in styles:
        styles["Heading 2"].font.name = "Calibri"
        styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        styles["Heading 2"].font.size = Pt(14)
        styles["Heading 2"].font.bold = True


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_footer_with_page_number(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    add_page_number(p)


def add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_title_block(doc: Document, title: str, subtitle: str = "", report_date: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)

    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(2)
        srun = sp.add_run(subtitle)
        srun.font.name = "Calibri"
        srun.font.size = Pt(12)
        srun.italic = True
        srun.font.color.rgb = RGBColor(89, 89, 89)

    if report_date:
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp.paragraph_format.space_after = Pt(10)
        drun = dp.add_run(report_date)
        drun.font.name = "Calibri"
        drun.font.size = Pt(11)
        drun.font.color.rgb = RGBColor(89, 89, 89)

    add_divider(doc)


def add_section_heading(doc: Document, text: str) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)


def add_body_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    if not text or not text.strip():
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text.strip())
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.bold = bold


def add_bullets_from_text(doc: Document, text: str) -> None:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        clean_line = line
        if clean_line.startswith("- "):
            clean_line = clean_line[2:]
        elif clean_line.startswith("• "):
            clean_line = clean_line[2:]

        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(clean_line)
        run.font.name = "Calibri"
        run.font.size = Pt(14)


def add_callout_box(doc: Document, heading: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F6FA")
    tc_pr.append(shd)

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(heading)
    r1.font.name = "Calibri"
    r1.font.size = Pt(13)
    r1.bold = True
    r1.font.color.rgb = RGBColor(31, 56, 100)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.1
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(12)

    doc.add_paragraph("")


def add_chart(doc: Document, image_path: str, caption: str = "", width_inches: float = 6.5) -> None:
    image_file = Path(image_path)
    if not image_file.exists():
        print(f"[WARN] Chart not found: {image_path}")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

    run = p.add_run()
    run.add_picture(str(image_file), width=Inches(width_inches))

    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        crun = cp.add_run(caption)
        crun.font.name = "Calibri"
        crun.font.size = Pt(11)
        crun.italic = True
        crun.font.color.rgb = RGBColor(89, 89, 89)


def render_section(doc: Document, heading: str, content) -> None:
    if content is None:
        return

    add_section_heading(doc, heading)

    if isinstance(content, str):
        paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)

    elif isinstance(content, list):
        for item in content:
            if isinstance(item, str):
                add_body_paragraph(doc, item)
            elif isinstance(item, dict):
                label = item.get("label", "")
                value = item.get("value", "")
                combined = f"{label}: {value}" if label else str(value)
                add_body_paragraph(doc, combined)

    elif isinstance(content, dict):
        for key, value in content.items():
            add_body_paragraph(doc, f"{key}: {value}")

    else:
        add_body_paragraph(doc, str(content))


def load_packet_json(packet_json_path: Path) -> dict:
    if not packet_json_path.exists():
        raise FileNotFoundError(
            f"Could not find packet JSON file at: {packet_json_path}. "
            f"Make sure build_weekly_market_packet.py writes output/weekly_market_packet.json"
        )

    with open(packet_json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def build_docx(packet_data: dict, output_path: Path) -> Path:
    doc = Document()
    set_document_defaults(doc)
    add_footer_with_page_number(doc)

    title = packet_data.get("title", "Weekly Market Packet")
    subtitle = packet_data.get("subtitle", "Institutional Market Commentary")
    report_date = packet_data.get("date", datetime.now().strftime("%B %d, %Y"))

    add_title_block(doc, title=title, subtitle=subtitle, report_date=report_date)

    executive_summary = packet_data.get("executive_summary", "")
    if executive_summary:
        add_callout_box(doc, "Executive Summary", executive_summary)

    section_order = [
        ("Market Overview", packet_data.get("market_overview", "")),
        ("Equity Market Trends", packet_data.get("equity_market_trends", "")),
        ("Rates and Macro Backdrop", packet_data.get("rates_and_macro_backdrop", "")),
        ("Institutional Signals", packet_data.get("institutional_signals", "")),
        ("Top Risks", packet_data.get("top_risks", "")),
        ("Closing Takeaways", packet_data.get("closing_takeaways", "")),
    ]

    for heading, content in section_order:
        if content:
            render_section(doc, heading, content)

        if heading == "Institutional Signals":
            charts = packet_data.get("charts", [])
            if charts:
                add_section_heading(doc, "Market Charts")
                for chart in charts:
                    if isinstance(chart, dict):
                        add_chart(
                            doc,
                            image_path=chart.get("path", ""),
                            caption=chart.get("caption", ""),
                            width_inches=chart.get("width_inches", 6.5),
                        )

    appendix_notes = packet_data.get("appendix_notes", "")
    if appendix_notes:
        render_section(doc, "Appendix / Notes", appendix_notes)

    doc.save(output_path)
    return output_path


def main() -> None:
    packet_data = load_packet_json(PACKET_JSON_PATH)
    output_file = build_docx(packet_data, DOCX_OUTPUT_PATH)
    print(f"[OK] DOCX created: {output_file}")


if __name__ == "__main__":
    main()
