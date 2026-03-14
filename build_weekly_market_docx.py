from pathlib import Path
import json
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

PACKET_JSON_PATH = OUTPUT_DIR / "weekly_market_packet.json"
DOCX_OUTPUT_PATH = OUTPUT_DIR / "weekly_market_packet.docx"


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(14)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_footer_with_page_number(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    add_page_number(p)


def add_title_block(doc: Document, title: str, subtitle: str, report_date: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sp.add_run(subtitle)
    srun.font.name = "Calibri"
    srun.font.size = Pt(12)
    srun.italic = True

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    drun = dp.add_run(report_date)
    drun.font.name = "Calibri"
    drun.font.size = Pt(11)

    doc.add_paragraph("")


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)


def add_body_paragraph(doc: Document, text: str) -> None:
    if not text or not str(text).strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text).strip())
    run.font.name = "Calibri"
    run.font.size = Pt(14)


def add_callout_box(doc: Document, heading: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    p1 = cell.paragraphs[0]
    r1 = p1.add_run(heading)
    r1.font.name = "Calibri"
    r1.font.size = Pt(13)
    r1.bold = True

    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(12)

    doc.add_paragraph("")


def add_chart(doc: Document, image_path: str, caption: str = "", width_inches: float = 6.5) -> None:
    image_file = Path(image_path)
    if not image_file.exists():
        print(f"[WARN] Chart not found: {image_path}")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_file), width=Inches(width_inches))

    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cp.add_run(caption)
        crun.font.name = "Calibri"
        crun.font.size = Pt(11)
        crun.italic = True


def render_section(doc: Document, heading: str, content: str) -> None:
    if not content:
        return
    add_section_heading(doc, heading)
    for paragraph in [p.strip() for p in str(content).split("\n") if p.strip()]:
        add_body_paragraph(doc, paragraph)


def load_packet_json() -> dict:
    if not PACKET_JSON_PATH.exists():
        raise FileNotFoundError(f"Could not find packet JSON file at: {PACKET_JSON_PATH}")
    with open(PACKET_JSON_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def build_docx(packet_data: dict) -> Path:
    doc = Document()
    set_document_defaults(doc)
    add_footer_with_page_number(doc)

    add_title_block(
        doc,
        title=packet_data.get("title", "Weekly Market Packet"),
        subtitle=packet_data.get("subtitle", "Institutional Market Commentary"),
        report_date=packet_data.get("date", datetime.now().strftime("%B %d, %Y")),
    )

    executive_summary = packet_data.get("executive_summary", "")
    if executive_summary:
        add_callout_box(doc, "Executive Summary", executive_summary)

    sections = [
        ("Market Overview", packet_data.get("market_overview", "")),
        ("Equity Market Trends", packet_data.get("equity_market_trends", "")),
        ("Rates and Macro Backdrop", packet_data.get("rates_and_macro_backdrop", "")),
        ("Institutional Signals", packet_data.get("institutional_signals", "")),
        ("Top Risks", packet_data.get("top_risks", "")),
        ("Closing Takeaways", packet_data.get("closing_takeaways", "")),
    ]

    for heading, content in sections:
        render_section(doc, heading, content)

        if heading == "Institutional Signals":
            charts = packet_data.get("charts", [])
            if charts:
                add_section_heading(doc, "Market Charts")
                for chart in charts:
                    add_chart(
                        doc,
                        image_path=chart.get("path", ""),
                        caption=chart.get("caption", ""),
                        width_inches=chart.get("width_inches", 6.5),
                    )

    appendix_notes = packet_data.get("appendix_notes", "")
    if appendix_notes:
        render_section(doc, "Appendix / Notes", appendix_notes)

    doc.save(DOCX_OUTPUT_PATH)

    if not DOCX_OUTPUT_PATH.exists():
        raise FileNotFoundError(f"DOCX was not created at expected path: {DOCX_OUTPUT_PATH}")

    return DOCX_OUTPUT_PATH


def main() -> None:
    print(f"[INFO] Loading packet JSON from: {PACKET_JSON_PATH}")
    packet_data = load_packet_json()
    output_file = build_docx(packet_data)
    print(f"[OK] DOCX created: {output_file.resolve()}")


if __name__ == "__main__":
    main()
